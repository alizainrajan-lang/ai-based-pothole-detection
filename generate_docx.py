import sys
import subprocess

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def set_cell_background(cell, color):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    cell._tc.get_or_add_tcPr().append(shading_elm)

doc = Document()

# Title
title = doc.add_heading('Software Test Plan', 0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Description Text
p = doc.add_paragraph()
p.add_run('Test Plan is the ').font.name = 'Calibri'
run = p.add_run('sub-set')
run.bold = True
run.font.name = 'Calibri'
p.add_run(' of the whole plan you had submitted earlier.\n').font.name = 'Calibri'
p.add_run('Although it has ').font.name = 'Calibri'
run2 = p.add_run('more details')
run2.bold = True
run2.font.name = 'Calibri'
p.add_run(' in terms of activities of the software testing your would do on your FYP.').font.name = 'Calibri'

p2 = doc.add_paragraph()
p2.add_run('For example if your software application has ').font.name = 'Calibri'
run3 = p2.add_run('4-Screens')
run3.bold = True
run3.font.name = 'Calibri'
p2.add_run(', ').font.name = 'Calibri'
run4 = p2.add_run('3-Reports')
run4.bold = True
run4.font.name = 'Calibri'
p2.add_run(' and ').font.name = 'Calibri'
run5 = p2.add_run('3-Modules')
run5.bold = True
run5.font.name = 'Calibri'
p2.add_run(' in it.\n').font.name = 'Calibri'
p2.add_run('Assume your estimation of efforts is approximately half day per screen / report').font.name = 'Calibri'

# Create table
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'

# Header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'S. No'
hdr_cells[1].text = 'Description'
hdr_cells[2].text = 'Test Engineer'
hdr_cells[3].text = 'Start Date'
hdr_cells[4].text = 'End Date'

for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.name = 'Calibri'

data = [
    # --- PHASE 1: INITIAL TESTING (JANUARY) ---
    ("1", "Main Dashboard UI Screen (Initial)", "Alizain Aslam", "05-Jan-2026", "06-Jan-2026"),
    ("2", "Live Scan Feed Sidebar Screen (Initial)", "Alizain Aslam", "08-Jan-2026", "09-Jan-2026"),
    ("3", "Interactive Map View Screen (Initial)", "Alizain Aslam", "12-Jan-2026", "14-Jan-2026"),
    ("4", "Location Search & Autocomplete Screen (Initial)\n(Note: Nominatim geocoding & autocomplete deferred/out of scope for current demo)", "Alizain Aslam", "19-Jan-2026", "20-Jan-2026"),
    ("", "- - -", "", "", ""),
    ("5", "Route Risk Assessment Report (Initial)", "Areeba Khan", "05-Jan-2026", "06-Jan-2026"),
    ("6", "Detection Summary Report (Initial)", "Areeba Khan", "08-Jan-2026", "09-Jan-2026"),
    ("7", "Pothole Detection Detail Report (Initial)\n(Note: Google Street View image integration deferred/out of scope due to billing hold)", "Areeba Khan", "12-Jan-2026", "13-Jan-2026"),
    ("", "- - -", "", "", ""),
    ("8", "Geolocation Tracking Module (Initial)\n(Note: Overpass road-snapping deferred/out of scope; markers sit at raw GPS coordinates)", "Alizain Aslam", "22-Jan-2026", "23-Jan-2026"),
    ("9", "External API Integrations Module (Initial)\n(Note: External API geocoding, road-snapping, and Street View integrations deferred/out of scope)", "Areeba Khan", "19-Jan-2026", "20-Jan-2026"),
    ("10", "Backend AI Engine API (Initial)", "Areeba Khan", "22-Jan-2026", "26-Jan-2026"), # Fixed Jan 24 (Sat) to Jan 26 (Mon)
    
    # --- PHASE 2 & 3: RETESTING & REGRESSION (MARCH) ---
    ("", "Bug Fix Buffer Phase (Development Team Fixing Issues — No Testing Activity)", "Dev Team", "27-Jan-2026", "27-Feb-2026"), # Highlighted yellow row separator
    
    ("11", "Retest: Main Dashboard & Live Feed", "Alizain Aslam", "02-Mar-2026", "03-Mar-2026"),
    ("12", "Retest: Interactive Map View", "Alizain Aslam", "04-Mar-2026", "05-Mar-2026"),
    ("13", "Retest: Location & Geolocation", "Alizain Aslam", "06-Mar-2026", "09-Mar-2026"), # Shifted to Mon
    ("", "- - -", "", "", ""),
    ("14", "Retest: All Analytical Reports", "Areeba Khan", "02-Mar-2026", "03-Mar-2026"),
    ("15", "Retest: External APIs & AI Engine", "Areeba Khan", "04-Mar-2026", "06-Mar-2026"),
    ("", "- - -", "", "", ""),
    ("16", "Full System Regression Testing", "Alizain & Areeba", "11-Mar-2026", "13-Mar-2026") # Fixed Mar 14 (Sat) to Mar 13 (Fri)
]

for idx, item in enumerate(data):
    row_cells = table.add_row().cells
    row_cells[0].text = item[0]
    row_cells[1].text = item[1]
    row_cells[2].text = item[2]
    row_cells[3].text = item[3]
    row_cells[4].text = item[4]
    
    # Make all cells slightly smaller to prevent page break issues
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
    
    # First item in red as per example
    if idx == 0:
        for paragraph in row_cells[1].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.bold = True
                
    # If it's a separator string "- - -", make it red
    if item[1] == "- - -":
        for paragraph in row_cells[1].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.bold = True
                
    # Bold the dates as in example
    if item[3]:
        for paragraph in row_cells[3].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    if item[4]:
        for paragraph in row_cells[4].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Yellow highlighted row
    if "Bug Fix Buffer Phase" in item[1]:
        for cell in row_cells:
            set_cell_background(cell, "FFFF00") # Yellow

# Adjust column widths
widths = (Inches(0.5), Inches(3.0), Inches(1.5), Inches(1.2), Inches(1.2))
for row in table.rows:
    for idx, width in enumerate(widths):
        row.cells[idx].width = width

doc.save('Software_Test_Plan_Final.docx')
print("Document generated successfully.")
